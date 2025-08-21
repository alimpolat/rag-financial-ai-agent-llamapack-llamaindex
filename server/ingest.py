import hashlib
import logging
from datetime import datetime
from pathlib import Path
from typing import List

import docx  # python-docx
from bs4 import BeautifulSoup
from llama_index.core import Document, StorageContext, load_index_from_storage
from pypdf import PdfReader

from .config import settings
from .llama_setup import configure_llama_index


def _extract_text(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        reader = PdfReader(str(file_path))
        parts = []
        for idx, page in enumerate(reader.pages):
            header = f"\n\n[Page {idx+1}]\n"
            parts.append(header + (page.extract_text() or ""))
        return "\n".join(parts)
    elif ext == ".docx":
        doc = docx.Document(str(file_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n\n".join(paragraphs)
    elif ext in {".html", ".htm"}:
        html = file_path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.extract()
        text = soup.get_text("\n")
        lines = [line.strip() for line in text.splitlines()]
        cleaned = "\n".join([l for l in lines if l])
        return cleaned
    elif ext in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _stable_doc_id(p: Path) -> str:
    h = hashlib.sha1()
    try:
        stat = p.stat()
        payload = f"{str(p.resolve())}|{stat.st_size}|{int(stat.st_mtime)}".encode("utf-8")
    except Exception:
        payload = str(p.resolve()).encode("utf-8")
    h.update(payload)
    return h.hexdigest()


def _to_documents(file_paths: List[Path]) -> List[Document]:
    docs: List[Document] = []
    ingested_at = datetime.utcnow().isoformat()
    for p in file_paths:
        ext = p.suffix.lower()
        base_meta = {
            "doc_id": _stable_doc_id(p),
            "source_path": str(p.resolve()),
            "file_name": p.name,
            "file_ext": ext,
            "file_size_bytes": p.stat().st_size,
            "ingested_at": ingested_at,
        }

        if ext == ".pdf":
            reader = PdfReader(str(p))
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if not page_text.strip():
                    continue
                metadata = {**base_meta, "page_number": idx + 1, "page_count": len(reader.pages)}
                docs.append(Document(text=page_text, metadata=metadata))
            continue

        if ext == ".docx":
            d = docx.Document(str(p))
            current_heading = None
            section_index = -1
            buffer: List[str] = []

            def flush_section():
                nonlocal buffer, section_index
                if not buffer:
                    return
                section_index += 1
                text = "\n\n".join(buffer).strip()
                if not text:
                    buffer = []
                    return
                metadata = {**base_meta, "section_index": section_index}
                if current_heading:
                    metadata["section_heading"] = current_heading
                docs.append(Document(text=text, metadata=metadata))
                buffer = []

            for para in d.paragraphs:
                content = (para.text or "").strip()
                if not content:
                    continue
                style_name = getattr(getattr(para, "style", None), "name", "") or ""
                if style_name.lower().startswith("heading"):
                    # Flush previous section and start a new one under this heading
                    flush_section()
                    current_heading = content
                    continue
                buffer.append(content)

            # Flush any remaining content
            flush_section()

            # Fallback: if no sections created, index whole doc
            if not any(
                doc.metadata.get("section_index") is not None
                for doc in docs
                if doc.metadata.get("doc_id") == base_meta["doc_id"]
            ):
                full_text = _extract_text(p)
                if full_text.strip():
                    docs.append(Document(text=full_text, metadata=base_meta))
            continue

        # txt/md or other supported types -> single document
        text = _extract_text(p)
        if text.strip():
            docs.append(Document(text=text, metadata=base_meta))
    return docs


def ingest_files(file_paths: List[str]) -> int:
    logger = logging.getLogger("ingest")
    configure_llama_index()
    # Normalize and de-duplicate input paths within a single request
    raw_paths = [Path(p) for p in file_paths]
    seen_realpaths = set()
    paths: List[Path] = []
    for p in raw_paths:
        try:
            rp = p.resolve()
        except Exception:
            rp = p
        if str(rp) in seen_realpaths:
            continue
        seen_realpaths.add(str(rp))
        paths.append(p)
    # Ensure storage dirs
    settings.storage_root.mkdir(parents=True, exist_ok=True)
    settings.upload_dir.mkdir(parents=True, exist_ok=True)
    settings.index_dir.mkdir(parents=True, exist_ok=True)

    documents = _to_documents(paths)
    logger.info("Prepared %d document(s) for ingestion", len(documents))

    # Try to load an existing index; if not present, initialize a new one
    try:
        storage_context = StorageContext.from_defaults(persist_dir=str(settings.index_dir))
        index = load_index_from_storage(storage_context=storage_context)

        # Collect existing doc_ids to skip duplicates
        existing_doc_ids = set()
        try:
            ds = index.storage_context.docstore
            for node in getattr(ds, "docs", {}).values():
                meta = getattr(node, "metadata", {}) or {}
                did = meta.get("doc_id")
                if did:
                    existing_doc_ids.add(did)
        except Exception:
            pass

        new_docs = [d for d in documents if d.metadata.get("doc_id") not in existing_doc_ids]
        if not new_docs:
            return 0
        for doc in new_docs:
            index.insert(doc)
        index.storage_context.persist(str(settings.index_dir))
        return len(new_docs)
    except Exception as e:
        logger.info("Initializing new index due to: %s", getattr(e, "__class__", type(e)).__name__)
        from llama_index.core import VectorStoreIndex

        storage_context = StorageContext.from_defaults()
        index = VectorStoreIndex.from_documents(documents, storage_context=storage_context)
        index.storage_context.persist(str(settings.index_dir))
        return len(documents)
